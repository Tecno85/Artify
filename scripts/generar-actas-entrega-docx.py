#!/usr/bin/env python3
"""Genera la versión Word de GA10-220501097-AA13-EV01."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs/proyecto/plan-capacitacion-actas-entrega-artify.docx"
FONT = "Times New Roman"
INK = "000000"
HEADER_FILL = "E7EEF8"
LIGHT_FILL = "F3F5F7"
TABLE_WIDTH = 9360


def set_run(run, size=12, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell(cell, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run(run, size=size, bold=bold)


def add_table(doc, headers, rows, widths, center_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
        set_cell_shading(table.rows[0].cells[index], HEADER_FILL)
        style_cell(table.rows[0].cells[index], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            align = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            style_cell(cells[index], align=align)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, *, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        set_run(paragraph.add_run(bold_lead), bold=True)
        set_run(paragraph.add_run(text[len(bold_lead):]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_plain(doc, text="", *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=4, size=12):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run(paragraph.add_run(text), size=size, bold=bold, italic=italic)
    return paragraph


def add_bullets(doc, items, checkbox=False):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run(paragraph.add_run(("☐ " if checkbox else "") + item))


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run(paragraph.add_run(item))


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = False
    run = paragraph.add_run(text)
    set_run(run, size=14 if level == 1 else 12, bold=True)
    return paragraph


def add_lines(doc, count=2):
    for _ in range(count):
        add_plain(doc, "________________________________________________________________________________", size=10, after=8)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr_text, fld_char2))
    set_run(run, size=12)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    add_page_number(section.header.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    for name, size, before, after in (("Heading 1", 14, 12, 6), ("Heading 2", 12, 10, 4), ("Heading 3", 12, 8, 3)):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def build_document():
    doc = Document()
    configure_document(doc)

    # Portada académica: variante APA del patrón editorial_cover.
    for _ in range(5):
        add_plain(doc, "", after=12)
    add_plain(doc, "Plan de capacitación y actas de entrega del proyecto Artify", bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=16, after=8)
    add_plain(doc, "Evidencia GA10-220501097-AA13-EV01", bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    for text in (
        "Iván Darío Madrid Daza",
        "Análisis y Desarrollo de Software",
        "Servicio Nacional de Aprendizaje (SENA)",
        "Instructor: José Ignacio Botero Osorio",
        "Julio de 2026",
    ):
        add_plain(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    doc.add_page_break()

    add_heading(doc, "Control del documento", 1)
    add_table(doc, ("Elemento", "Descripción"), (
        ("Documento", "Plan de capacitación y modelos de actas de Artify"),
        ("Versión", "1.0"),
        ("Fecha de elaboración", "Julio de 2026"),
        ("Autor y capacitador", "Iván Darío Madrid Daza"),
        ("Producto", "Artify — editor de imágenes web"),
        ("Estado", "Modelo listo para programar y diligenciar"),
    ), (2700, 6660))
    note = add_plain(doc, "Nota de diligenciamiento. Este documento contiene el plan y los modelos solicitados. Las fechas de ejecución, los datos de los participantes, las respuestas y las firmas se completan cuando ocurran la capacitación y la entrega.", italic=True, size=10, after=8)
    note.paragraph_format.left_indent = Inches(0.25)
    note.paragraph_format.right_indent = Inches(0.25)

    add_heading(doc, "1. Introducción", 1)
    add_body(doc, "En esta evidencia presento el plan con el que capacitaré a las personas que usarán Artify, el acta para registrar su satisfacción y el acta con la que formalizaré la entrega del software. Organicé los instrumentos en el orden en que se aplican durante la implantación: capacitación, comprobación del aprendizaje y entrega del producto.")
    add_body(doc, "La entrega debe involucrar a las partes interesadas y dejar información verificable sobre el producto, la aceptación y las responsabilidades posteriores. Este enfoque es coherente con ISO/IEC/IEEE 12207:2026. Los criterios de aceptación también permiten valorar el cumplimiento del propósito del producto, de acuerdo con ISO/IEC 25010:2023.")

    doc.add_page_break()
    add_heading(doc, "2. Objetivos", 1)
    add_heading(doc, "2.1 Objetivo general", 2)
    add_body(doc, "Planificar la capacitación y formalizar la satisfacción y la entrega de Artify mediante instrumentos claros, verificables y adecuados para las personas que usarán o recibirán el sistema.")
    add_heading(doc, "2.2 Objetivos específicos", 2)
    add_bullets(doc, (
        "Explicar el acceso y las funciones principales de Artify.",
        "Comprobar un flujo básico de edición y descarga.",
        "Registrar la satisfacción, las dudas y las observaciones.",
        "Identificar los componentes entregados y dejar constancia de la conformidad.",
    ))

    add_heading(doc, "3. Plan de capacitación", 1)
    add_heading(doc, "3.1 Datos generales", 2)
    add_table(doc, ("Elemento", "Información"), (
        ("Nombre", "Capacitación para el uso de Artify"),
        ("Responsable", "Iván Darío Madrid Daza"),
        ("Participantes", "Usuario operativo y, cuando corresponda, administrador"),
        ("Modalidad", "Demostración guiada y práctica supervisada"),
        ("Lugar y fecha", "Por diligenciar"),
        ("Duración", "60 minutos; 75 minutos con rol administrador"),
        ("Recursos", "Equipo con navegador, Artify, imágenes de prueba y manuales"),
    ), (2500, 6860))
    add_heading(doc, "3.2 Resultados esperados", 2)
    add_bullets(doc, (
        "Ingresar a Artify y reconocer su interfaz.",
        "Cargar una imagen y aplicar una transformación y un filtro.",
        "Deshacer o rehacer un cambio y descargar el resultado.",
        "Cerrar la sesión de forma segura.",
        "Gestionar cuentas, si el participante tiene rol administrador.",
    ))
    add_heading(doc, "3.3 Agenda", 2)
    add_table(doc, ("Momento", "Actividad", "Tiempo"), (
        ("Inicio", "Propósito y comprobación de acceso", "5 min"),
        ("Reconocimiento", "Acceso e identificación de la interfaz", "5 min"),
        ("Editor", "Carga, transformaciones, filtros e historial", "20 min"),
        ("Resultado", "Conversión, descarga, perfil y cierre", "10 min"),
        ("Práctica", "Ejercicio completo del participante", "15 min"),
        ("Cierre", "Preguntas y evaluación", "5 min"),
        ("Total capacitación operativa", "", "60 min"),
        ("Extensión administrativa opcional", "Gestión controlada de usuarios", "15 min adicionales"),
        ("Total con rol administrador", "", "75 min"),
    ), (1700, 6060, 1600), center_columns=(2,))
    add_plain(doc, "La capacitación operativa dura 60 minutos. Si participa una persona con rol administrador, agrego un bloque de 15 minutos, para un total de 75 minutos.", italic=True, size=10)
    add_heading(doc, "3.4 Metodología y evaluación", 2)
    add_body(doc, "Explicaré cada operación, la demostraré y pediré al participante que repita el flujo con una imagen de prueba sin datos confidenciales. Utilizaré las guías de usuario como apoyo.")
    add_numbered(doc, (
        "Iniciar sesión.", "Cargar una imagen válida.", "Aplicar y confirmar una herramienta.",
        "Deshacer o rehacer un cambio.", "Descargar el resultado.", "Cerrar la sesión.",
    ))
    add_plain(doc, "Resultado: ☐ Cumplido   ☐ Cumplido con apoyo   ☐ Pendiente de refuerzo", bold=True)

    add_heading(doc, "4. Acta de satisfacción de la capacitación", 1)
    add_heading(doc, "4.1 Identificación", 2)
    add_table(doc, ("Elemento", "Información"), (
        ("Código", "ART-CAP-____"), ("Proyecto", "Artify"),
        ("Lugar o medio", "________________________________________"),
        ("Fecha", "____ / ____ / ______"), ("Horario", "__________ a __________"),
        ("Capacitador", "Iván Darío Madrid Daza"),
        ("Persona capacitada", "________________________________________"),
        ("Identificación o cargo", "________________________________________"),
        ("Rol", "☐ Usuario operativo   ☐ Administrador"),
    ), (2700, 6660))
    add_heading(doc, "4.2 Desarrollo y resultado", 2)
    add_plain(doc, "Objetivo: orientar al participante en el uso correcto y seguro de las funciones de Artify correspondientes a su rol.")
    add_plain(doc, "Temas desarrollados", bold=True)
    add_bullets(doc, (
        "Acceso e interfaz", "Carga y edición de imágenes", "Historial, conversión y descarga",
        "Preferencias, perfil y cierre seguro", "Gestión administrativa de usuarios",
        "Manuales y atención de dudas",
    ), checkbox=True)
    add_plain(doc, "Resultado práctico: ☐ Cumplido   ☐ Cumplido con apoyo   ☐ Pendiente de refuerzo", bold=True)

    add_heading(doc, "4.3 Valoración del participante", 2)
    add_plain(doc, "Marque una opción: 1 = muy insatisfecho, 2 = insatisfecho, 3 = aceptable, 4 = satisfecho y 5 = muy satisfecho.", size=10)
    add_table(doc, ("Criterio", "1", "2", "3", "4", "5"), (
        ("Claridad de la explicación", "☐", "☐", "☐", "☐", "☐"),
        ("Utilidad de los temas", "☐", "☐", "☐", "☐", "☐"),
        ("Dominio del capacitador", "☐", "☐", "☐", "☐", "☐"),
        ("Tiempo y ritmo", "☐", "☐", "☐", "☐", "☐"),
        ("Capacidad para usar Artify", "☐", "☐", "☐", "☐", "☐"),
    ), (4860, 900, 900, 900, 900, 900), center_columns=(1, 2, 3, 4, 5))
    add_plain(doc, "¿Se cumplieron los objetivos?   ☐ Sí   ☐ Parcialmente   ☐ No", bold=True)
    add_plain(doc, "Conclusiones, observaciones o temas por reforzar:", bold=True)
    add_lines(doc, 2)
    add_plain(doc, "Compromisos y fecha de seguimiento, si aplica:", bold=True)
    add_lines(doc, 1)
    doc.add_page_break()
    add_heading(doc, "4.4 Declaración y firmas", 2)
    add_plain(doc, "Declaro que recibí la capacitación sobre Artify, tuve la oportunidad de practicar, formular preguntas y registrar mis observaciones. Mi firma acredita la participación y la información consignada; no elimina los compromisos pendientes.", size=11)
    add_table(doc, ("Persona capacitada", "Capacitador"), (
        ("Firma: __________________________", "Firma: __________________________"),
        ("Nombre: _________________________", "Iván Darío Madrid Daza"),
        ("Fecha: __________________________", "Fecha: __________________________"),
    ), (4680, 4680))

    add_heading(doc, "5. Acta final de entrega del software", 1)
    add_heading(doc, "5.1 Identificación de las partes", 2)
    add_table(doc, ("Elemento", "Información"), (
        ("Código", "ART-ENT-____"),
        ("Lugar y fecha", "____________________, ____ / ____ / ______"),
        ("Persona que entrega", "Iván Darío Madrid Daza"),
        ("Persona que recibe", "________________________________________"),
        ("Identificación", "________________________________________"),
        ("Cargo o entidad", "________________________________________"),
        ("Medio de contacto", "________________________________________"),
    ), (2700, 6660))
    add_heading(doc, "5.2 Producto entregado", 2)
    add_body(doc, "Yo, Iván Darío Madrid Daza, hago entrega de Artify, una aplicación web para cargar, transformar y descargar imágenes. El producto incluye un frontend con HTML, CSS y JavaScript, una API con Node.js y Express, y persistencia en PostgreSQL.")
    add_body(doc, "La versión entregada permite autenticarse, editar imágenes, utilizar el historial, convertir y descargar archivos, guardar preferencias y consultar actividad. También incluye un panel protegido para gestionar cuentas de usuario.")
    doc.add_page_break()
    add_heading(doc, "5.3 Relación de entregables", 2)
    add_table(doc, ("Entregable", "Ubicación", "Verificado"), (
        ("Aplicación web", "tecno85.github.io/artify/", "☐"),
        ("Código fuente", "github.com/Tecno85/artify", "☐"),
        ("Guías de usuario", "docs/proyecto/", "☐"),
        ("Manual técnico", "docs/tecnica/manual-tecnico-artify.md", "☐"),
        ("Instalación y despliegue", "docs/tecnica/", "☐"),
        ("Artefactos PostgreSQL", "database/postgresql/", "☐"),
        ("Mantenimiento y respaldo", "docs/tecnica/", "☐"),
    ), (3300, 4860, 1200), center_columns=(2,))
    add_plain(doc, "Seguridad: las credenciales, secretos, tokens y cadenas privadas de conexión no forman parte de los anexos públicos. Todo acceso autorizado debe transferirse por un canal seguro.", italic=True, size=10)

    add_heading(doc, "5.4 Verificación y aceptación", 2)
    add_bullets(doc, (
        "Acceso a la aplicación y al repositorio acordado.",
        "Inicio de sesión con una cuenta autorizada.",
        "Carga, edición y descarga de una imagen de prueba.",
        "Acceso a manuales y documentos.",
        "Registro de novedades, limitaciones o compromisos.",
        "Capacitación realizada o fecha acordada para programarla.",
    ), checkbox=True)
    add_plain(doc, "Estado: ☐ Recibida a satisfacción   ☐ Recibida con observaciones   ☐ No aceptada", bold=True)
    add_plain(doc, "Observaciones, pendientes y fechas acordadas:", bold=True)
    add_lines(doc, 4)
    add_plain(doc, "Condiciones o periodo de soporte acordado:", bold=True)
    add_lines(doc, 2)
    doc.add_page_break()
    add_heading(doc, "5.5 Declaración y firmas", 2)
    add_plain(doc, "La persona que recibe declara que pudo verificar los elementos marcados y que conoce las observaciones, limitaciones y compromisos escritos. La firma formaliza la entrega, pero no supone aceptar elementos expresamente pendientes.", size=11)
    add_table(doc, ("Persona que recibe", "Persona que entrega"), (
        ("Firma: __________________________", "Firma: __________________________"),
        ("Nombre: _________________________", "Iván Darío Madrid Daza"),
        ("Identificación: _________________", "Identificación: _________________"),
        ("Fecha: __________________________", "Fecha: __________________________"),
    ), (4680, 4680))

    add_heading(doc, "6. Conclusiones", 1)
    add_body(doc, "Con este plan puedo orientar la capacitación hacia las tareas que el usuario realmente necesita y comprobar el aprendizaje mediante un ejercicio breve. El acta de satisfacción registra la percepción del participante sin confundirla con la entrega definitiva.")
    add_body(doc, "El acta final reúne las partes, el alcance de Artify, los componentes entregados, las verificaciones, las observaciones y las firmas. Así dejo una evidencia sencilla y trazable sin afirmar que una actividad se realizó antes de contar con sus datos.")
    add_heading(doc, "Referencias", 1)
    references = (
        "International Organization for Standardization. (2023). ISO/IEC 25010:2023: Systems and software engineering—Systems and software Quality Requirements and Evaluation (SQuaRE)—Product quality model. https://www.iso.org/standard/78176.html",
        "International Organization for Standardization. (2026). ISO/IEC/IEEE 12207:2026: Systems and software engineering—Software life cycle processes. https://www.iso.org/standard/90219.html",
        "Madrid Daza, I. D. (2026a). Guía del usuario administrador de Artify [Documento del proyecto]. Repositorio Artify. https://github.com/Tecno85/artify/blob/main/docs/proyecto/guia-usuario-administrador-artify.md",
        "Madrid Daza, I. D. (2026b). Guía del usuario operativo de Artify [Documento del proyecto]. Repositorio Artify. https://github.com/Tecno85/artify/blob/main/docs/proyecto/guia-usuario-operativo-artify.md",
    )
    for reference in references:
        paragraph = add_plain(doc, reference, size=11, after=8)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.line_spacing = 1.5

    doc.core_properties.title = "Plan de capacitación y actas de entrega del proyecto Artify"
    doc.core_properties.subject = "Evidencia GA10-220501097-AA13-EV01"
    doc.core_properties.author = "Iván Darío Madrid Daza"
    doc.core_properties.keywords = "Artify, capacitación, acta de satisfacción, entrega de software, SENA"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
